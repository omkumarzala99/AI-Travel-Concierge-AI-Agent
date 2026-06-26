import os
import io
from typing import List
from langchain_core.documents import Document
import pypdf
import docx
from app.logger import get_logger, log_execution_time

log = get_logger("rag.loaders")

@log_execution_time("rag.loaders")
def load_document_from_bytes(file_bytes: bytes, file_name: str) -> List[Document]:
    """
    Parses document byte streams (PDF, TXT, DOCX) and returns a list of LangChain Document objects.
    """
    _, ext = os.path.splitext(file_name.lower())
    documents = []

    try:
        if ext == ".txt":
            # Direct text decoding
            text = file_bytes.decode("utf-8", errors="ignore")
            documents.append(Document(
                page_content=text,
                metadata={"source": file_name}
            ))
            
        elif ext == ".pdf":
            # PDF byte parsing using PyPDF
            pdf_file = io.BytesIO(file_bytes)
            reader = pypdf.PdfReader(pdf_file)
            for page_num, page in enumerate(reader.pages):
                text = page.extract_text() or ""
                if text.strip():
                    documents.append(Document(
                        page_content=text,
                        metadata={
                            "source": file_name,
                            "page": page_num + 1,
                            "total_pages": len(reader.pages)
                        }
                    ))
                    
        elif ext == ".docx":
            # Word file parsing using python-docx
            docx_file = io.BytesIO(file_bytes)
            doc = docx.Document(docx_file)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            full_text = "\n\n".join(paragraphs)
            if full_text.strip():
                documents.append(Document(
                    page_content=full_text,
                    metadata={"source": file_name}
                ))
        else:
            log.warning(f"Unsupported file format: {ext} for file '{file_name}'")
            
    except Exception as e:
        log.error(f"Error parsing file '{file_name}': {str(e)}", exc_info=True)
        raise ValueError(f"Failed to parse document '{file_name}': {str(e)}")

    return documents
